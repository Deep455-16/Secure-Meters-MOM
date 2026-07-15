"""
MOM Generator REST API Server
Provides HTTP endpoints for MOM generation
Run on http://127.0.0.1:5001
"""

import os
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"

# Load .env file if present (for any machine running this app)
def _load_dotenv():
    env_path = os.path.join(os.path.dirname(__file__), '.env')
    if os.path.exists(env_path):
        with open(env_path, encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith('#') and '=' in line:
                    key, _, value = line.partition('=')
                    os.environ.setdefault(key.strip(), value.strip())
_load_dotenv()

# Always ensure HuggingFace symlink crash is prevented on Windows
os.environ.setdefault("HF_HUB_DISABLE_SYMLINKS", "1")
os.environ.setdefault("HF_HUB_DISABLE_SYMLINKS_WARNING", "1")

# Add source code/backend to Python path
import sys
_BASE_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, os.path.join(_BASE_DIR, 'source code', 'backend'))

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from flask_socketio import SocketIO, emit
from stream_handler import VADAudioStream
from transcriber_streaming import StreamingTranscriber
from streaming_mom import StreamingMOMPipeline
import os
import json
import logging
import io
from docx import Document
from fpdf import FPDF
from datetime import datetime
from flask import send_from_directory
from mom_generator import MOMGenerator

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Initialize Flask app
app = Flask(__name__)
CORS(app)
socketio = SocketIO(app, cors_allowed_origins="*", async_mode="threading")

# Global streaming instances
streaming_transcriber = None
streaming_mom_pipeline = None
vad_streams = {}

@socketio.on('connect')
def handle_connect():
    sid = request.sid
    logger.info(f"Client connected for streaming: {sid}")
    vad_streams[sid] = VADAudioStream(sample_rate=16000)
    emit('status', {'message': 'Connected to audio stream'})

@socketio.on('disconnect')
def handle_disconnect():
    sid = request.sid
    logger.info(f"Client disconnected: {sid}")
    if sid in vad_streams:
        del vad_streams[sid]

@socketio.on('audio_chunk')
def handle_audio_chunk(data):
    sid = request.sid
    if sid not in vad_streams:
        return
        
    vad = vad_streams[sid]
    segments = vad.process_chunk(data)
    
    if streaming_transcriber and segments:
        for seg in segments:
            streaming_transcriber.add_audio_segment(seg)

def transcription_callback(text):
    socketio.emit('transcript_update', {'text': text})
    print(f"\n[🎤 LIVE ASR Transcribed]: {text}\n", flush=True)
    if streaming_mom_pipeline:
        streaming_mom_pipeline.add_transcript_text(text)

@app.route('/api/v1/generate-live-mom', methods=['POST'])
def generate_live_mom():
    if not streaming_mom_pipeline:
        return jsonify({'error': 'Streaming pipeline not initialized'}), 500
        
    try:
        def emit_progress(pct, status_text):
            socketio.emit('progress_update', {'progress': pct, 'status': status_text})
        
        # Stop stream immediately — discard pending audio, don't process it
        emit_progress(10, "Stream stopped. Discarding pending audio...")
        if streaming_transcriber:
            streaming_transcriber.stop()
        streaming_mom_pipeline.stop()

        # Early bail-out: if no insights AND no raw transcript were collected, skip the LLM entirely
        if not streaming_mom_pipeline.memory_store.get_all().strip() and not streaming_mom_pipeline.raw_store.get_all().strip():
            emit_progress(100, "No meeting content detected.")
            streaming_mom_pipeline.start()
            if streaming_transcriber:
                streaming_transcriber.start()
            return jsonify({'status': 'ok', 'mom': 'No meeting content was captured during this stream. Please ensure your microphone is working and try again.'}), 200

        # Generate the final MOM from collected insights
        emit_progress(50, "Generating final Live MOM from collected insights...")
        
        data = request.get_json(silent=True) or {}
        duration = data.get('duration')
        
        final_mom = streaming_mom_pipeline.generate_final_mom(duration=duration)
        
        # Restart the pipeline and transcriber for the next session
        streaming_mom_pipeline.start()
        if streaming_transcriber:
            streaming_transcriber.start()
        
        emit_progress(100, "Live MOM Generation Complete!")
        return jsonify({'status': 'ok', 'mom': final_mom}), 200
    except Exception as e:
        logger.error(f"Live MOM generation failed: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/v1/export/docx', methods=['POST'])
def export_docx():
    data = request.json
    text = data.get('text', '')
    doc = Document()
    doc.add_heading('Minutes of Meeting', 0)
    for line in text.split('\n'):
        doc.add_paragraph(line)
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return send_file(file_stream, as_attachment=True, download_name='MOM.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/api/v1/export/pdf', methods=['POST'])
def export_pdf():
    data = request.json
    text = data.get('text', '')
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    
    for line in text.split('\n'):
        clean_line = line.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 8, txt=clean_line)
        
    # fpdf2: output() returns bytes directly — wrap in BytesIO for send_file
    pdf_bytes = pdf.output()
    file_stream = io.BytesIO(bytes(pdf_bytes))
    file_stream.seek(0)
    return send_file(file_stream, as_attachment=True, download_name='MOM.pdf', mimetype='application/pdf')

@app.route('/api/v1/export/template', methods=['POST'])
def export_template():
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor, Inches, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import re, copy

        data = request.json
        text = data.get('text', '')
        if not text:
            return jsonify({'error': 'No MOM text provided'}), 400

        # ── Regex helpers ──────────────────────────────────────────────────
        def extract_section(pattern, src_, default=''):
            m = re.search(pattern, src_, re.IGNORECASE | re.DOTALL)
            return m.group(1).strip() if m else default

        def extract_list_from_block(block):
            lines = []
            for ln in block.splitlines():
                ln = re.sub(r'\*{1,2}', '', ln)
                ln = re.sub(r'^[\s\-\*\d\.\)]+', '', ln).strip()
                if ln:
                    lines.append(ln)
            return lines

        def extract_list(pattern, src_):
            block = extract_section(pattern, src_)
            return extract_list_from_block(block) if block else []

        NS = (r'(?=\n\s*\n?\s*\*{0,2}\d*\.?\s*'
              r'(?:DATE|TIME|ATTENDEES?|PARTICIPANTS?|TOPICS? DISCUSSED|AGENDA|'
              r'KEY DISCUSSION|DECISIONS?(?: MADE| TAKEN)?|ACTION ITEMS?|'
              r'NEXT STEPS?|MINUTES OF MEETING)\b|\Z)')

        # Metadata (anchored to line-start)
        meeting_date = extract_section(r'(?:^|\n)\*{0,2}(?:DATE|MEETING DATE)\*{0,2}\s*[:\-]\s*([^\n*]+)', text)
        meeting_time = extract_section(r'(?:^|\n)\*{0,2}TIME\*{0,2}\s*[:\-]\s*([^\n*]+)', text)
        purpose      = extract_section(r'(?:^|\n)\*{0,2}(?:PURPOSE|SUBJECT|MEETING PURPOSE)\*{0,2}\s*[:\-]\s*([^\n*]+)', text)
        participants = extract_section(r'\*{0,2}(?:ATTENDEES?|PARTICIPANTS?|MEMBERS? PRESENT|PRESENT)\*{0,2}\s*[:\-]\s*(.*?)' + NS, text)
        duration     = extract_section(r'(?:^|\n)\*{0,2}(?:DURATION|MEETING DURATION|TOTAL DURATION|LENGTH)\*{0,2}\s*[:\-]\s*([^\n*]+)', text)

        if not duration:
            dur_m = re.search(r'(?:meeting lasted|duration of|ran for|total time)[:\s]+([^\n\.]+)', text, re.IGNORECASE)
            duration = dur_m.group(1).strip() if dur_m else ''

        agenda_items = extract_list(r'\*{0,2}(?:\d+\.\s*)?(?:AGENDA\s*/?\s*)?TOPICS?\s*DISCUSSED?\*{0,2}\s*[:\-]?\s*(.*?)' + NS, text)
        if not agenda_items:
            agenda_items = extract_list(r'\*{0,2}(?:\d+\.\s*)?AGENDA\*{0,2}\s*[:\-]?\s*(.*?)' + NS, text)

        if not purpose:
            title_m = re.search(r'MINUTES\s+OF\s+MEETING\s*[-\u2013:]\s*([^\n]+)', text, re.IGNORECASE)
            if title_m:
                purpose = title_m.group(1).strip()
            elif agenda_items:
                purpose = '; '.join(agenda_items[:2])

        discussion_items = extract_list(r'\*{0,2}(?:\d+\.\s*)?KEY\s*DISCUSSION\s*POINTS?\*{0,2}\s*[:\-]?\s*(.*?)' + NS, text)
        if not discussion_items:
            discussion_items = extract_list(r'\*{0,2}(?:\d+\.\s*)?(?:DISCUSSION(?: SUMMARY)?|KEY POINTS?|HIGHLIGHTS?)\*{0,2}\s*[:\-]?\s*(.*?)' + NS, text)

        decision_items = extract_list(r'\*{0,2}(?:\d+\.\s*)?DECISIONS?\s*(?:MADE|TAKEN)?\*{0,2}\s*[:\-]?\s*(.*?)' + NS, text)
        if not decision_items:
            decision_items = extract_list(r'\*{0,2}(?:\d+\.\s*)?RESOLUTIONS?\*{0,2}\s*[:\-]?\s*(.*?)' + NS, text)

        action_raw     = extract_section(r'\*{0,2}(?:\d+\.\s*)?ACTION\s*ITEMS?\*{0,2}\s*[:\-]?\s*(.*?)' + NS, text)
        next_steps_raw = extract_section(r'\*{0,2}(?:\d+\.\s*)?NEXT\s*STEPS?\*{0,2}\s*[:\-]?\s*(.*?)' + NS, text)

        def parse_actions(raw_block, label_prefix=''):
            items = []
            for ln in extract_list_from_block(raw_block):
                if not ln:
                    continue
                resp_m = re.search(r'\b(?:owner|responsible|by)\s*[:\-]?\s*([A-Z][a-zA-Z\s]+)', ln, re.IGNORECASE)
                date_m = re.search(r'\b(\d{1,2}[\s\-/]\w+[\s\-/]?\d{0,4}|ASAP|immediately)\b', ln, re.IGNORECASE)
                items.append({
                    'desc': f"{label_prefix}{ln}" if label_prefix else ln,
                    'resp': resp_m.group(1).strip() if resp_m else '',
                    'date': date_m.group(1) if date_m else ''
                })
            return items

        actions = parse_actions(action_raw)
        actions.extend(parse_actions(next_steps_raw, label_prefix='[Next Step] '))

        # ── Load template as layout/border base ───────────────────────────
        template_path = os.path.join(_BASE_DIR, 'source code', 'backend', 'SECURE_MOM_Template.docx')
        doc = DocxDocument(template_path)
        WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

        # ── XML helpers ───────────────────────────────────────────────────
        def _run_el(text_val, bold=False, font='Calibri', size_pt=10, color=None):
            """Build a <w:r> element with explicit formatting."""
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rf = OxmlElement('w:rFonts')
            rf.set(qn('w:ascii'), font); rf.set(qn('w:hAnsi'), font)
            rPr.append(rf)
            if bold:
                rPr.append(OxmlElement('w:b'))
                rPr.append(OxmlElement('w:bCs'))
            if color:
                c = OxmlElement('w:color'); c.set(qn('w:val'), color); rPr.append(c)
            sz = OxmlElement('w:sz');   sz.set(qn('w:val'),   str(int(size_pt * 2)));  rPr.append(sz)
            szc = OxmlElement('w:szCs'); szc.set(qn('w:val'), str(int(size_pt * 2))); rPr.append(szc)
            r.append(rPr)
            t = OxmlElement('w:t')
            t.text = str(text_val)
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            r.append(t)
            return r

        def _shd_el(hex_color):
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), hex_color)
            return shd

        def write_cell(cell, value, bold=False, font='Calibri', size_pt=10,
                       color=None, bg=None, align='left'):
            tc = cell._tc
            # Add/update shading
            tcPr = tc.find(f'{{{WNS}}}tcPr')
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr'); tc.insert(0, tcPr)
            if bg:
                for old in tcPr.findall(f'{{{WNS}}}shd'): tcPr.remove(old)
                tcPr.append(_shd_el(bg))
            # Remove all existing paragraphs
            for p in tc.findall(f'{{{WNS}}}p'): tc.remove(p)
            # Build paragraph
            p = OxmlElement('w:p')
            pPr = OxmlElement('w:pPr')
            jc = OxmlElement('w:jc'); jc.set(qn('w:val'), align); pPr.append(jc)
            p.append(pPr)
            for li, line in enumerate(str(value).split('\n')):
                p.append(_run_el(line, bold=bold, font=font, size_pt=size_pt, color=color))
                if li < len(str(value).split('\n')) - 1:
                    br_r = OxmlElement('w:r'); br = OxmlElement('w:br'); br_r.append(br); p.append(br_r)
            tc.append(p)

        def make_row(ref_row, values, bold=False, font='Calibri', size_pt=10, bg=None):
            """Clone ref_row's cell widths/borders, fill each cell with one clean run."""
            new_tr = copy.deepcopy(ref_row._tr)
            cells  = new_tr.findall(f'{{{WNS}}}tc')
            for ci, cell_el in enumerate(cells):
                val = str(values[ci]) if ci < len(values) else ''
                if bg:
                    tcPr = cell_el.find(f'{{{WNS}}}tcPr')
                    if tcPr is None:
                        tcPr = OxmlElement('w:tcPr'); cell_el.insert(0, tcPr)
                    for old in tcPr.findall(f'{{{WNS}}}shd'): tcPr.remove(old)
                    tcPr.append(_shd_el(bg))
                for p in cell_el.findall(f'{{{WNS}}}p'): cell_el.remove(p)
                p = OxmlElement('w:p')
                p.append(_run_el(val, bold=bold, font=font, size_pt=size_pt))
                cell_el.append(p)
            return new_tr

        def fill_dynamic_table(tbl, has_header, rows_data,
                               header_labels=None, font='Calibri', size_pt=10):
            tbl_el = tbl._tbl
            if has_header:
                # Row 0 = proper header: apply bold+grey explicitly
                hdr_tr = tbl.rows[0]._tr
                for cell_el in hdr_tr.findall(f'{{{WNS}}}tc'):
                    tcPr = cell_el.find(f'{{{WNS}}}tcPr')
                    if tcPr is None:
                        tcPr = OxmlElement('w:tcPr'); cell_el.insert(0, tcPr)
                    for old in tcPr.findall(f'{{{WNS}}}shd'): tcPr.remove(old)
                    tcPr.append(_shd_el('D9D9D9'))
                    for r_el in cell_el.findall(f'.//{{{WNS}}}r'):
                        rPr = r_el.find(f'{{{WNS}}}rPr')
                        if rPr is None:
                            rPr = OxmlElement('w:rPr'); r_el.insert(0, rPr)
                        if rPr.find(f'{{{WNS}}}b') is None:
                            rPr.insert(0, OxmlElement('w:b'))
                        if rPr.find(f'{{{WNS}}}bCs') is None:
                            rPr.append(OxmlElement('w:bCs'))
                # Row 1 = Jinja template row: use for cloning then remove
                ref_row = tbl.rows[1]
                tbl_el.remove(ref_row._tr)
            else:
                # Both rows are Jinja markers
                ref_row = tbl.rows[1]
                for row in list(tbl.rows):
                    tbl_el.remove(row._tr)
                # Add header row (cloned from ref_row structure, filled with labels)
                if header_labels:
                    hdr_tr = make_row(ref_row, header_labels,
                                      bold=True, font=font, size_pt=size_pt, bg='D9D9D9')
                    tbl_el.append(hdr_tr)

            # Add data rows
            for row_vals in rows_data:
                tbl_el.append(make_row(ref_row, row_vals,
                                       bold=False, font=font, size_pt=size_pt))
            if not rows_data:
                n = len(list(ref_row._tr.findall(f'{{{WNS}}}tc')))
                tbl_el.append(make_row(ref_row, [''] * n,
                                       bold=False, font=font, size_pt=size_pt))

        # ════════════════════════════════════════════════════════════════
        # TABLE 0 — Header
        # ════════════════════════════════════════════════════════════════
        ht = doc.tables[0]

        # ── Top-left cell: insert the SECURE logo image ───────────────
        logo_cell = ht.cell(0, 0)
        tc = logo_cell._tc
        # Clear existing paragraphs from the cell
        for p in tc.findall(f'{{{WNS}}}p'):
            tc.remove(p)
        # Build a new centred paragraph and add the image run
        logo_p = OxmlElement('w:p')
        logo_pPr = OxmlElement('w:pPr')
        logo_jc = OxmlElement('w:jc')
        logo_jc.set(qn('w:val'), 'center')
        logo_pPr.append(logo_jc)
        logo_p.append(logo_pPr)
        tc.append(logo_p)

        logo_path = os.path.join(_BASE_DIR, 'source code', 'assets', 'secure_logo.png')
        if os.path.exists(logo_path):
            # Insert image into the cell paragraph
            logo_para = logo_cell.paragraphs[0]
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = logo_para.add_run()
            run.add_picture(logo_path, width=Inches(1.3))
        else:
            # Fallback: plain text if logo file missing
            write_cell(logo_cell, 'SECURE', bold=True, size_pt=14,
                       color='1F3864', align='center')
        write_cell(ht.cell(0,1), 'GENERAL\nMinutes Of Meeting',
                   bold=True, size_pt=10, align='center')
        write_cell(ht.cell(0,2),
                   f'Doc Date   {meeting_date}\nPage       1 of 1',
                   bold=False, size_pt=9, align='left')
        write_cell(ht.cell(1,0), '', bold=False, size_pt=10)
        write_cell(ht.cell(1,1), 'Minutes Of Meeting',
                   bold=True, size_pt=11, align='center')
        write_cell(ht.cell(1,2), '', bold=False, size_pt=10)

        # ════════════════════════════════════════════════════════════════
        # TABLE 1 — Info
        # ════════════════════════════════════════════════════════════════
        it = doc.tables[1]
        rows_info = [
            (0, 'Meeting Date', meeting_date, 'Time', meeting_time),
            (1, 'Venue', '', '', ''),
            (2, 'Purpose', purpose, '', ''),
            (3, 'Planned Duration', '', 'Actual Duration', duration),
            (4, 'Meeting Effort\n(in person hrs.)', '', 'Convened By', ''),
            (5, 'Minutes Taken By', 'AI Generator', 'Reviewed By', ''),
        ]
        for ri, lbl1, val1, lbl2, val2 in rows_info:
            write_cell(it.cell(ri,0), lbl1, bold=True, size_pt=9, bg='D9D9D9')
            write_cell(it.cell(ri,1), val1, bold=False, size_pt=9)
            if lbl2:
                write_cell(it.cell(ri,2), lbl2, bold=True, size_pt=9, bg='D9D9D9')
                write_cell(it.cell(ri,3), val2, bold=False, size_pt=9)

        # ════════════════════════════════════════════════════════════════
        # TABLE 2 — Participants
        # ════════════════════════════════════════════════════════════════
        pt = doc.tables[2]
        for ci, hd in enumerate(['Members Present', 'Members Absent', 'Invitees']):
            write_cell(pt.cell(0,ci), hd, bold=True, size_pt=9,
                       bg='D9D9D9', align='center')
        write_cell(pt.cell(1,0), participants, bold=False, size_pt=9)
        write_cell(pt.cell(1,1), '', bold=False, size_pt=9)
        write_cell(pt.cell(1,2), '', bold=False, size_pt=9)

        # ════════════════════════════════════════════════════════════════
        # TABLE 3 — Agenda (both rows are Jinja markers)
        # ════════════════════════════════════════════════════════════════
        fill_dynamic_table(doc.tables[3], has_header=False,
            rows_data=[[f'{i+1}.', item] for i,item in enumerate(agenda_items)],
            header_labels=['No.', 'Agenda Item'], size_pt=9)

        # ════════════════════════════════════════════════════════════════
        # TABLE 4 — Decisions
        # ════════════════════════════════════════════════════════════════
        fill_dynamic_table(doc.tables[4], has_header=True,
            rows_data=[[str(i+1), item] for i,item in enumerate(decision_items)],
            size_pt=9)

        # ════════════════════════════════════════════════════════════════
        # TABLE 5 — Discussion Summary
        # ════════════════════════════════════════════════════════════════
        fill_dynamic_table(doc.tables[5], has_header=True,
            rows_data=[[str(i+1), item] for i,item in enumerate(discussion_items)],
            size_pt=9)

        # ════════════════════════════════════════════════════════════════
        # TABLE 6 — Action Plan
        # ════════════════════════════════════════════════════════════════
        fill_dynamic_table(doc.tables[6], has_header=True,
            rows_data=[[str(i+1), a.get('desc',''), a.get('resp',''),
                        a.get('date',''), ''] for i,a in enumerate(actions)],
            size_pt=9)

        # ── Save & serve ──────────────────────────────────────────────
        os.makedirs('moms', exist_ok=True)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename  = f'SECURE_MOM_Template_{timestamp}.docx'
        save_path = os.path.join('moms', filename)
        doc.save(save_path)
        logger.info(f"MOM Template saved to {save_path}")
        return send_file(save_path, as_attachment=True,
                         download_name=filename,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    except Exception as e:
        logger.error(f"Template export failed: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500



# Initialize MOM Generator (deferred to background thread)
generator = None


# Store for tracking active jobs
jobs = {}
job_counter = 0


@app.route('/health', methods=['GET'])
def health():
    """Health check endpoint"""
    try:
        # Check if ollama is running
        ollama_ok = False
        try:
            models = generator.ollama.list_models()
            ollama_ok = len(models) > 0
        except Exception:
            ollama_ok = False
        
        return jsonify({
            'status': 'healthy',
            'timestamp': datetime.now().isoformat(),
            'ollama_running': ollama_ok,
            'generator_initialized': generator is not None
        }), 200
    except Exception as e:
        return jsonify({
            'status': 'unhealthy',
            'error': str(e)
        }), 500


@app.route('/api/v1/status', methods=['GET'])
def status():
    """Get system status"""
    try:
        status_info = {
            'version': '1.0',
            'timestamp': datetime.now().isoformat(),
            'services': {}
        }
        
        # Check Ollama
        if generator:
            try:
                models = generator.ollama.list_models()
                status_info['services']['ollama'] = {
                    'status': 'running',
                    'url': generator.config.get('ollama_url'),
                    'models': models[:5]  # First 5 models
                }
            except Exception:
                status_info['services']['ollama'] = {'status': 'offline'}
            
            # Check Whisper
            try:
                import faster_whisper  # noqa: F401
                whisper_ready = True
            except ImportError:
                whisper_ready = False
                
            status_info['services']['whisper'] = {
                'status': 'ready' if whisper_ready else 'offline',
                'executable': whisper_ready,
                'model': whisper_ready
            }
        
        return jsonify(status_info), 200
    except Exception as e:
        logger.error(f"Status check error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/v1/models', methods=['GET'])
def get_models():
    """Get available Ollama models"""
    try:
        if not generator:
            return jsonify({'error': 'Generator not initialized'}), 500
        
        models = generator.ollama.list_models()
        return jsonify({
            'models': models,
            'current_model': generator.config.get('ollama_model')
        }), 200
    except Exception as e:
        logger.error(f"Models error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/v1/templates', methods=['GET'])
def get_templates():
    """Get available MOM templates"""
    try:
        from mom_templates import MOMTemplateManager
        templates = MOMTemplateManager.load_templates()
        
        template_list = []
        for name, template_info in templates.items():
            template_list.append({
                'name': name,
                'title': template_info.get('name'),
                'description': template_info.get('description')
            })
        
        return jsonify({'templates': template_list}), 200
    except Exception as e:
        logger.error(f"Templates error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/v1/generate', methods=['POST'])
def generate_mom():
    """Generate MOM from audio file"""
    try:
        if not generator:
            return jsonify({'error': 'Generator not initialized'}), 500
        
        # Check if audio file is provided
        if 'audio' not in request.files:
            return jsonify({'error': 'No audio file provided'}), 400
        
        audio_file = request.files['audio']
        if audio_file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Get template if provided
        template = request.form.get('template', 'standard')
        
        # Get word limit if provided
        word_limit = request.form.get('word_limit', type=int, default=300)
        
        # Save uploaded file temporarily
        temp_path = os.path.join('recordings', audio_file.filename)
        os.makedirs('recordings', exist_ok=True)
        audio_file.save(temp_path)
        
        logger.info(f"Processing: {audio_file.filename}")
        
        try:
            # Process the meeting with progress callback
            def emit_progress(pct, status_text):
                socketio.emit('progress_update', {'progress': pct, 'status': status_text})
                
            result = generator.process_meeting(temp_path, template=template, word_limit=word_limit, progress_callback=emit_progress)
            # Read generated MOM
            with open(result['mom_file'], 'r', encoding='utf-8') as f:
                mom_content = f.read()
            
            # Read transcript
            with open(result['transcript_file'], 'r', encoding='utf-8') as f:
                transcript_content = f.read()
            
            response = {
                'status': 'success',
                'filename': audio_file.filename,
                'template': template,
                'transcript': transcript_content,
                'mom': mom_content,
                'files': {
                    'transcript': result['transcript_file'],
                    'mom': result['mom_file']
                }
            }
            
            logger.info(f"Successfully processed: {audio_file.filename}")
            return jsonify(response), 200
        
        finally:
            # Clean up temp file
            if os.path.exists(temp_path):
                os.remove(temp_path)
    
    except Exception as e:
        logger.error(f"Generation error: {str(e)}")
        return jsonify({'error': str(e), 'type': type(e).__name__}), 500


@app.route('/api/v1/transcribe', methods=['POST'])
def transcribe_only():
    """Transcribe audio without generating MOM"""
    try:
        if not generator:
            return jsonify({'error': 'Generator not initialized'}), 500
        
        if 'audio' not in request.files:
            return jsonify({'error': 'No audio file provided'}), 400
        
        audio_file = request.files['audio']
        if audio_file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Save temporarily
        temp_path = os.path.join('recordings', audio_file.filename)
        os.makedirs('recordings', exist_ok=True)
        audio_file.save(temp_path)
        
        logger.info(f"Transcribing: {audio_file.filename}")
        
        try:
            # Transcribe only
            transcript = generator.transcriber.transcribe(temp_path)
            
            return jsonify({
                'status': 'success',
                'filename': audio_file.filename,
                'transcript': transcript
            }), 200
        
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)
    
    except Exception as e:
        logger.error(f"Transcription error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/v1/generate-from-text', methods=['POST'])
def generate_mom_from_text():
    """Generate MOM from text transcript directly"""
    try:
        if not generator:
            return jsonify({'error': 'Generator not initialized'}), 500
        
        data = request.get_json()
        if not data or 'transcript' not in data:
            return jsonify({'error': 'No transcript provided'}), 400
        
        transcript = data['transcript']
        template = data.get('template', 'standard')
        word_limit = data.get('word_limit', 300)
        
        logger.info(f"Generating MOM from text transcript (template: {template}, word limit: {word_limit})")
        
        try:
            # Generate MOM
            mom = generator.ollama.generate_mom(transcript, generator.config.get('ollama_model', 'llama2'), template, word_limit)
            
            return jsonify({
                'status': 'success',
                'template': template,
                'transcript': transcript,
                'mom': mom
            }), 200
        
        except Exception as e:
            logger.error(f"MOM generation error: {str(e)}")
            return jsonify({'error': str(e)}), 500
    
    except Exception as e:
        logger.error(f"Request error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/v1/download/<file_type>/<filename>', methods=['GET'])
def download_file(file_type, filename):
    """Download transcript or MOM file"""
    try:
        if file_type == 'transcript':
            filepath = os.path.join('transcripts', filename)
        elif file_type == 'mom':
            filepath = os.path.join('moms', filename)
        else:
            return jsonify({'error': 'Invalid file type'}), 400
        
        if not os.path.exists(filepath):
            return jsonify({'error': 'File not found'}), 404
        
        return send_file(filepath, as_attachment=True, mimetype='text/plain')
    
    except Exception as e:
        logger.error(f"Download error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/v1/files', methods=['GET'])
def list_files():
    """List all generated files"""
    try:
        files_info = {
            'transcripts': [],
            'moms': []
        }
        
        # List transcripts
        if os.path.exists('transcripts'):
            for f in os.listdir('transcripts'):
                if f.endswith('.txt'):
                    fpath = os.path.join('transcripts', f)
                    files_info['transcripts'].append({
                        'name': f,
                        'size': os.path.getsize(fpath),
                        'modified': datetime.fromtimestamp(os.path.getmtime(fpath)).isoformat()
                    })
        
        # List MOMs
        if os.path.exists('moms'):
            for f in os.listdir('moms'):
                if f.endswith('.txt'):
                    fpath = os.path.join('moms', f)
                    files_info['moms'].append({
                        'name': f,
                        'size': os.path.getsize(fpath),
                        'modified': datetime.fromtimestamp(os.path.getmtime(fpath)).isoformat()
                    })
        
        return jsonify(files_info), 200
    
    except Exception as e:
        logger.error(f"List files error: {str(e)}")
        return jsonify({'error': str(e)}), 500


# ============== TEAMS INTEGRATION ENDPOINTS ==============

@app.route('/api/v1/teams/status', methods=['GET'])
def teams_status():
    """Get Teams integration status"""
    try:
        # Try to import Teams integration module
        try:
            from teams_integration import configure_teams_integration
            teams_manager = configure_teams_integration("config.json")
            
            if teams_manager and teams_manager.is_configured():
                status = teams_manager.get_system_status()
            else:
                status = {
                    "teams_configured": False,
                    "message": "Teams credentials not configured in config.json",
                    "setup_required": True
                }
        except ImportError:
            status = {
                "teams_configured": False,
                "message": "Teams integration module not available",
                "setup_required": False
            }
        
        return jsonify({
            'status': 'ok',
            'teams': status,
            'timestamp': datetime.now().isoformat()
        }), 200
    
    except Exception as e:
        logger.error(f"Teams status error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/v1/teams/connect', methods=['POST'])
def teams_connect():
    """Initialize connection to Microsoft Teams"""
    try:
        data = request.get_json()
        
        # Validate required fields
        required = ['client_id', 'client_secret']
        if not all(field in data for field in required):
            return jsonify({'error': 'Missing required fields: client_id, client_secret'}), 400
        
        from teams_integration import TeamsAuthManager
        
        auth_manager = TeamsAuthManager(
            client_id=data['client_id'],
            client_secret=data['client_secret'],
            tenant_id=data.get('tenant_id', 'common')
        )
        
        # Test connection
        token = auth_manager.get_token()
        if not token:
            return jsonify({'error': 'Failed to authenticate with Teams API'}), 401
        
        logger.info("Teams connection established successfully")
        return jsonify({
            'status': 'connected',
            'message': 'Successfully connected to Microsoft Teams',
            'timestamp': datetime.now().isoformat()
        }), 200
    
    except Exception as e:
        logger.error(f"Teams connect error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/v1/teams/list-calls', methods=['GET'])
def teams_list_calls():
    """List active Teams calls"""
    try:
        from teams_integration import configure_teams_integration
        teams_manager = configure_teams_integration("config.json")
        
        if not teams_manager or not teams_manager.is_configured():
            return jsonify({'error': 'Teams not configured'}), 400
        
        if not teams_manager.call_recorder:
            return jsonify({'error': 'Call recorder not available'}), 500
        
        calls = teams_manager.call_recorder.list_active_calls()
        
        return jsonify({
            'status': 'success',
            'calls': calls,
            'count': len(calls),
            'timestamp': datetime.now().isoformat()
        }), 200
    
    except Exception as e:
        logger.error(f"List calls error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/v1/teams/participants/<meeting_id>', methods=['GET'])
def teams_get_participants(meeting_id):
    """Get participants in a Teams meeting"""
    try:
        from teams_integration import configure_teams_integration
        teams_manager = configure_teams_integration("config.json")
        
        if not teams_manager or not teams_manager.is_configured():
            return jsonify({'error': 'Teams not configured'}), 400
        
        if not teams_manager.call_recorder:
            return jsonify({'error': 'Call recorder not available'}), 500
        
        participants = teams_manager.call_recorder.get_meeting_participants(meeting_id)
        
        return jsonify({
            'status': 'success',
            'meeting_id': meeting_id,
            'participants': participants,
            'count': len(participants),
            'timestamp': datetime.now().isoformat()
        }), 200
    
    except Exception as e:
        logger.error(f"Get participants error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/v1/teams/generate-mom/<meeting_id>', methods=['POST'])
def teams_generate_mom(meeting_id):
    """Generate MOM from Teams meeting transcript"""
    try:
        from teams_integration import configure_teams_integration
        teams_manager = configure_teams_integration("config.json")
        
        if not teams_manager or not teams_manager.is_configured():
            return jsonify({'error': 'Teams not configured'}), 400
        
        if not teams_manager.transcript_processor:
            return jsonify({'error': 'Transcript processor not available'}), 500
        
        data = request.get_json() or {}
        template = data.get('template', 'standard')
        word_limit = data.get('word_limit', 300)
        
        logger.info(f"Processing Teams meeting {meeting_id} with template: {template}")
        
        # Process the Teams meeting
        result = teams_manager.process_teams_meeting(meeting_id, 'moms')
        
        if not result:
            return jsonify({'error': 'Failed to process Teams meeting. Transcript may not be available yet.'}), 400
        
        # Generate MOM from transcript
        try:
            mom = generator.ollama.generate_mom(result['transcript'], generator.config.get('ollama_model', 'llama2'), template, word_limit)
            
            return jsonify({
                'status': 'success',
                'meeting_id': meeting_id,
                'participants': result['participants'],
                'transcript': result['transcript'],
                'mom': mom,
                'template': template,
                'timestamp': datetime.now().isoformat()
            }), 200
        
        except Exception as e:
            logger.error(f"MOM generation error: {str(e)}")
            return jsonify({'error': f'MOM generation failed: {str(e)}'}), 500
    
    except Exception as e:
        logger.error(f"Teams MOM generation error: {str(e)}")
        return jsonify({'error': str(e)}), 500


# ============================================================

@app.route('/api/v1/teams/test', methods=['POST'])
def teams_test_connection():
    """Test MS Teams API Credentials"""
    try:
        data = request.get_json()
        if not data or 'api_key' not in data or 'tenant_id' not in data or 'client_id' not in data:
            return jsonify({'error': 'Missing Teams credentials'}), 400
        
        # Test connection (placeholder logic assuming api_key handles auth checking)
        from teams_integration import TeamsAuthManager
        auth = TeamsAuthManager(client_id=data['client_id'], client_secret=data['api_key'], tenant_id=data['tenant_id'])
        if auth.get_token():
            return jsonify({'status': 'success', 'message': 'Connected to MS Teams successfully!'}), 200
        else:
            return jsonify({'error': 'Authentication failed. Please check credentials.'}), 401
    except Exception as e:
        logger.error(f"Teams test connection error: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/v1/teams/fetch-meeting', methods=['POST'])
def teams_fetch_meeting():
    """Fetch meeting from MS Teams Link"""
    try:
        def emit_progress(pct, status_text):
            socketio.emit('progress_update', {'progress': pct, 'status': status_text})
            
        emit_progress(10, "Authenticating with MS Teams API...")
        
        data = request.get_json()
        meeting_link = data.get('meeting_link')
        
        if not meeting_link:
            return jsonify({'error': 'Missing meeting link'}), 400
            
        # Parse meeting_id from link (very simplified, usually graph API needs proper ID)
        meeting_id = meeting_link.split('/')[-1] if '/' in meeting_link else meeting_link
        
        # Initialize and fetch using manager
        from teams_integration import TeamsIntegrationManager
        config = {
            "teams_client_id": data.get('client_id'),
            "teams_client_secret": data.get('api_key'),
            "teams_tenant_id": data.get('tenant_id', 'common')
        }
        manager = TeamsIntegrationManager(config)
        
        if not manager.is_configured():
            return jsonify({'error': 'Failed to configure Teams manager.'}), 400
            
        emit_progress(30, "Fetching transcript from MS Teams...")
        result = manager.process_teams_meeting(meeting_link, 'moms')
        if result:
            emit_progress(60, "Transcript fetched successfully. Setting up generator...")
            # Generate MOM directly
            template = data.get('template', 'standard')
            word_limit = data.get('word_limit', 300)
            try:
                # We assume generator is globally available as it is defined at module level
                mom_model = generator.config.get('ollama_model', 'qwen2.5:7b')
                emit_progress(80, f"Generating MOM using {mom_model}...")
                
                mom = generator.ollama.generate_mom(result['transcript'], mom_model, template, word_limit)
                
                emit_progress(100, "MOM Generation Complete!")
                
                # Save MOM
                os.makedirs('moms', exist_ok=True)
                mom_file = os.path.join('moms', f"{meeting_id}_teams_mom.txt")
                with open(mom_file, 'w', encoding='utf-8') as f:
                    f.write(mom)
                    
                return jsonify({
                    'status': 'success',
                    'meeting_id': meeting_id,
                    'transcript': result['transcript'],
                    'participants': result['participants'],
                    'mom': mom
                }), 200
            except Exception as e:
                logger.error(f"Failed to generate MOM after fetching Teams meeting: {str(e)}")
                return jsonify({
                    'status': 'success',
                    'meeting_id': meeting_id,
                    'transcript': result['transcript'],
                    'participants': result['participants'],
                    'warning': 'Fetched successfully but MOM generation failed.'
                }), 200
        else:
            return jsonify({'error': 'Could not fetch or process meeting. Check credentials or link.'}), 400
    except Exception as e:
        logger.error(f"Teams fetch meeting error: {str(e)}")
        return jsonify({'error': str(e)}), 500


# ── Frontend page routes ──────────────────────────────────────────────────

FRONTEND_DIR = os.path.join(_BASE_DIR, 'source code', 'frontend')


def _serve_html(filename):
    """Helper: read and serve an HTML file from the frontend directory."""
    path = os.path.join(FRONTEND_DIR, filename)
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return f.read(), 200, {'Content-Type': 'text/html; charset=utf-8'}
    except FileNotFoundError:
        return f"{filename} not found — make sure the frontend files are in 'source code/frontend/'", 404


@app.route('/', methods=['GET'])
def home():
    """Main MOM Generator dashboard."""
    return _serve_html('index.html')


@app.route('/live', methods=['GET'])
@app.route('/live-stream', methods=['GET'])
def live_stream_page():
    """Live Stream ASR feature page."""
    return _serve_html('live_stream.html')


@app.route('/upload', methods=['GET'])
@app.route('/recording', methods=['GET'])
def upload_recording_page():
    """Upload Recording feature page."""
    return _serve_html('upload_recording.html')


@app.route('/transcript', methods=['GET'])
@app.route('/text', methods=['GET'])
def text_transcript_page():
    """Text Transcript feature page."""
    return _serve_html('text_transcript.html')


@app.route('/teams', methods=['GET'])
def teams_page():
    """Microsoft Teams Integration feature page."""
    return _serve_html('teams_integration.html')


@app.route('/settings', methods=['GET'])
def settings_page():
    """Settings page."""
    return _serve_html('mom_settings.html')


@app.route('/socket.io.min.js')
def serve_socket_io():
    """Serve local socket.io file."""
    return send_from_directory(FRONTEND_DIR, 'socket.io.min.js')




@app.route('/api/v1/settings', methods=['GET'])
def get_settings():
    """Get current user settings"""
    try:
        # Settings are stored in browser localStorage, but this endpoint
        # can return server-side settings if needed
        return jsonify({
            'status': 'ok',
            'server_settings': {
                'max_audio_upload': '500MB',
                'max_video_upload': '1GB',
                'supported_formats': ['wav', 'mp3', 'm4a', 'mp4', 'ogg', 'aac', 'flac'],
                'supported_video': ['mp4', 'mkv', 'avi']
            },
            'timestamp': datetime.now().isoformat()
        }), 200
    except Exception as e:
        logger.error(f"Settings error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/v1/settings/validate', methods=['POST'])
def validate_settings():
    """Validate user settings"""
    try:
        data = request.get_json()
        
        # Validate audio length
        audio_max = data.get('audioMaxLength', 30)
        if not (1 <= audio_max <= 180):
            return jsonify({'valid': False, 'error': 'Audio max length must be 1-180 minutes'}), 400
        
        # Validate text length
        text_max = data.get('textMaxLength', 10000)
        if not (100 <= text_max <= 1000000):
            return jsonify({'valid': False, 'error': 'Text max length must be 100-1,000,000 characters'}), 400
        
        return jsonify({
            'valid': True,
            'message': 'Settings are valid',
            'timestamp': datetime.now().isoformat()
        }), 200
    
    except Exception as e:
        logger.error(f"Validation error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/v1/config', methods=['GET'])
def get_config():
    """Get current configuration from config.json"""
    try:
        if generator and hasattr(generator, 'config'):
            return jsonify({
                'status': 'ok',
                'config': generator.config,
                'timestamp': datetime.now().isoformat()
            }), 200
        else:
            return jsonify({'status': 'error', 'message': 'Generator not initialized'}), 500
    except Exception as e:
        logger.error(f"Config retrieval error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/v1/config', methods=['POST'])
def save_config():
    """Save configuration updates to config.json"""
    try:
        data = request.get_json()
        
        # Validate required fields
        if not data:
            return jsonify({'status': 'error', 'message': 'No configuration data provided'}), 400
        
        # Load existing config
        config_file = os.path.join(_BASE_DIR, "config.json")
        with open(config_file, 'r') as f:
            config = json.load(f)
        
        # Update with new values (only update provided keys)
        for key in ['chunking_model', 'ollama_model', 'ollama_url', 'whisper_exe', 'whisper_model', 
                    'output_dir', 'recordings_dir', 'transcripts_dir', 'moms_dir']:
            if key in data:
                config[key] = data[key]
        
        # Save updated config back to file
        with open(config_file, 'w') as f:
            json.dump(config, f, indent=4)
            
        # Update running in-memory instances
        if generator:
            generator.config = config
            if hasattr(generator, 'ollama'):
                generator.ollama.base_url = config.get('ollama_url', 'http://localhost:11434').rstrip('/')
                
        if streaming_mom_pipeline:
            streaming_mom_pipeline.chunking_model = config.get('chunking_model', 'qwen2.5:1.5b')
            streaming_mom_pipeline.refinement_model = config.get('ollama_model', 'qwen2.5:7b')
        
        return jsonify({
            'status': 'ok',
            'message': 'Configuration saved successfully',
            'config': config,
            'timestamp': datetime.now().isoformat()
        }), 200
        
    except FileNotFoundError:
        return jsonify({'status': 'error', 'message': 'config.json file not found'}), 404
    except json.JSONDecodeError:
        return jsonify({'status': 'error', 'message': 'Invalid JSON in config file'}), 400
    except Exception as e:
        logger.error(f"Config save error: {str(e)}")
        return jsonify({'error': str(e)}), 500


# ============================================================


@app.route('/', methods=['GET'])
def index():
    """Web UI - Simple HTML interface"""
    html = """
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <title>MOM Generator | Whisper + Ollama</title>
        <meta name="viewport" content="width=device-width, initial-scale=1">
        <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;600;700&display=swap" rel="stylesheet">
        <!-- Socket.IO served locally for full offline support -->
        <script src="/socket.io.min.js"></script>
        <style>
            :root {
                --primary: #4F46E5;
                --primary-hover: #4338CA;
                --bg: #F3F4F6;
                --bg-gradient: linear-gradient(135deg, #fdfbfb 0%, #ebedee 100%);
                --card-bg: rgba(255, 255, 255, 0.9);
                --text: #1F2937;
                --text-light: #6B7280;
                --border: #E5E7EB;
                --success: #10B981;
                --error: #EF4444;
                --input-bg: #F9FAFB;
                --result-bg: #F8FAFC;
            }
            
            /* Dark Mode Variables */
            [data-theme="dark"] {
                --primary: #6366F1;
                --primary-hover: #4F46E5;
                --bg: #111827;
                --bg-gradient: linear-gradient(135deg, #111827 0%, #1F2937 100%);
                --card-bg: rgba(31, 41, 55, 0.9);
                --text: #F9FAFB;
                --text-light: #9CA3AF;
                --border: #374151;
                --input-bg: #374151;
                --result-bg: #1F2937;
            }

            body {
                font-family: 'Inter', sans-serif;
                background: var(--bg-gradient);
                color: var(--text);
                margin: 0;
                padding: 40px 20px;
                min-height: 100vh;
                transition: background 0.3s, color 0.3s;
            }
            .container {
                max-width: 800px;
                margin: 0 auto;
                background: var(--card-bg);
                backdrop-filter: blur(10px);
                border-radius: 16px;
                padding: 40px;
                box-shadow: 0 10px 25px rgba(0,0,0,0.1);
                border: 1px solid rgba(128,128,128,0.2);
                transition: background 0.3s;
            }
            .header {
                display: flex;
                justify-content: space-between;
                align-items: flex-start;
                margin-bottom: 40px;
            }
            .header-text {
                text-align: left;
            }
            h1 {
                font-weight: 700;
                color: var(--primary);
                margin: 0 0 10px 0;
                font-size: 2.5rem;
            }
            p.subtitle {
                color: var(--text-light);
                font-size: 1.1rem;
                margin: 0;
            }
            
            /* Theme Toggle Button */
            .theme-toggle {
                background: var(--border);
                color: var(--text);
                border: none;
                padding: 8px 16px;
                border-radius: 20px;
                cursor: pointer;
                font-weight: 600;
                font-size: 0.9rem;
                transition: background 0.2s;
                width: auto;
                display: inline-block;
            }
            .theme-toggle:hover { background: var(--text-light); color: var(--bg); }

            .status-bar {
                display: flex;
                justify-content: flex-start;
                gap: 15px;
                margin-bottom: 30px;
            }
            .status-badge {
                padding: 6px 12px;
                border-radius: 20px;
                font-size: 0.85rem;
                font-weight: 600;
                display: flex;
                align-items: center;
                gap: 6px;
            }
            .status-ok { background: rgba(16, 185, 129, 0.15); color: var(--success); }
            .status-error { background: rgba(239, 68, 68, 0.15); color: var(--error); }
            
            .section {
                background: transparent;
                border: 1px solid var(--border);
                border-radius: 12px;
                padding: 25px;
                margin-bottom: 25px;
                transition: transform 0.2s, box-shadow 0.2s;
            }
            .section:hover {
                box-shadow: 0 4px 12px rgba(0,0,0,0.08);
            }
            .teams-section {
                background: rgba(79, 70, 229, 0.05);
                border-color: var(--primary);
            }
            h2 {
                font-size: 1.25rem;
                font-weight: 600;
                margin-top: 0;
                margin-bottom: 20px;
                display: flex;
                align-items: center;
                gap: 8px;
            }
            .form-group {
                margin-bottom: 20px;
            }
            label {
                display: block;
                font-size: 0.9rem;
                font-weight: 600;
                margin-bottom: 8px;
                color: var(--text);
            }
            input[type="file"], select, textarea {
                width: 100%;
                padding: 12px;
                border: 1px solid var(--border);
                border-radius: 8px;
                font-family: 'Inter', sans-serif;
                font-size: 0.95rem;
                box-sizing: border-box;
                background: var(--input-bg);
                color: var(--text);
                transition: border-color 0.2s, background 0.3s, color 0.3s;
            }
            select:focus, textarea:focus {
                outline: none;
                border-color: var(--primary);
                box-shadow: 0 0 0 3px rgba(79, 70, 229, 0.2);
            }
            textarea {
                resize: vertical;
                min-height: 120px;
            }
            button {
                background: var(--primary);
                color: white;
                border: none;
                padding: 12px 24px;
                border-radius: 8px;
                font-size: 1rem;
                font-weight: 600;
                cursor: pointer;
                width: 100%;
                transition: background 0.2s, transform 0.1s;
                display: flex;
                justify-content: center;
                align-items: center;
                gap: 8px;
            }
            button:hover { background: var(--primary-hover); }
            button:active { transform: scale(0.98); }
            button:disabled { background: var(--border); color: var(--text-light); cursor: not-allowed; transform: none; }
            
            button.danger { background: var(--error); }
            button.danger:hover { background: #DC2626; }
            
            button.secondary { background: transparent; border: 2px solid var(--primary); color: var(--primary); }
            button.secondary:hover { background: rgba(79, 70, 229, 0.1); }

            /* Recording Indicator */
            .rec-indicator {
                display: none;
                align-items: center;
                gap: 8px;
                color: var(--error);
                font-weight: bold;
                margin-bottom: 15px;
            }
            .rec-indicator.active { display: flex; }
            .rec-dot {
                width: 12px; height: 12px;
                background: var(--error);
                border-radius: 50%;
                animation: pulse 1s infinite;
            }
            @keyframes pulse { 0% { transform: scale(1); opacity: 1; } 50% { transform: scale(1.5); opacity: 0.5; } 100% { transform: scale(1); opacity: 1; } }
            
            .result {
                margin-top: 20px;
                background: var(--result-bg);
                border: 1px solid var(--border);
                border-radius: 8px;
                padding: 20px;
                font-size: 0.95rem;
                line-height: 1.6;
                white-space: pre-wrap;
                display: none;
                color: var(--text);
            }
            .result.active { display: block; }
            .result.error {
                background: rgba(239, 68, 68, 0.1);
                border-color: rgba(239, 68, 68, 0.3);
                color: var(--error);
            }
            
            /* Loading Spinner */
            .spinner {
                display: none;
                width: 20px;
                height: 20px;
                border: 3px solid rgba(255,255,255,0.3);
                border-radius: 50%;
                border-top-color: white;
                animation: spin 1s ease-in-out infinite;
            }
            @keyframes spin { to { transform: rotate(360deg); } }
            .btn-loading .spinner { display: inline-block; }
            .btn-loading .btn-text { display: none; }
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                <div class="header-text">
                    <h1>📋 Minutes Generator</h1>
                    <div id="modelConfigDisplay" style="font-size: 0.85rem; color: var(--primary); margin-top: 8px; font-weight: 600;"></div>
                    <p class="subtitle">AI-powered concise meeting summaries</p>
                </div>
                <div style="display: flex; gap: 10px;">
                    <button class="theme-toggle" onclick="toggleTheme()">🌗 Theme</button>
                    <button class="theme-toggle" onclick="window.location.href='/settings'">⚙️ Settings</button>
                </div>
            </div>
            
            <div id="status" class="status-bar"></div>
            
            <!-- MS Teams Integration Section -->
            <div class="section teams-section" id="teamsSection">
                <h2>🌐 MS Teams Live Integration</h2>
                <p style="margin-top:0; color:var(--text-light); font-size:0.9rem;">
                    Connect your Azure AD app to fetch Teams meeting transcripts and auto-generate MOMs.
                </p>

                <!-- Credentials Form -->
                <div id="teamsCredForm">
                    <div class="form-group">
                        <label for="teamsClientId">Azure AD Client ID</label>
                        <input type="text" id="teamsClientId" placeholder="xxxxxxxx-xxxx-xxxx-xxxx-xxxxxxxxxxxx" autocomplete="off">
                    </div>
                    <div class="form-group">
                        <label for="teamsClientSecret">Client Secret</label>
                        <input type="password" id="teamsClientSecret" placeholder="Your Azure AD client secret" autocomplete="off">
                    </div>
                    <div class="form-group">
                        <label for="teamsTenantId">Tenant ID</label>
                        <input type="text" id="teamsTenantId" placeholder="xxxxxxxx-xxxx-xxxx-xxxx-xxxxxxxxxxxx (or 'common')" autocomplete="off">
                    </div>
                    <div style="display:flex; gap:10px; margin-bottom:15px;">
                        <button id="btnTeamsTest" onclick="teamsTestConnection()" class="secondary">
                            🔌 Test Connection
                        </button>
                        <button id="btnTeamsConnect" onclick="teamsConnect()" style="flex:1;">
                            🔗 Connect to MS Teams
                        </button>
                    </div>
                    <!-- MS Teams Progress Bar -->
                    <div id="progressContainerTeams" class="progress-container" style="display:none; margin-bottom:15px;">
                        <div class="progress-bar-bg" style="width:100%; background:#333; border-radius:5px;">
                            <div id="progressBarFillTeams" style="width:0%; background:#4CAF50; height:15px; border-radius:5px; transition: width 0.3s;"></div>
                        </div>
                        <p id="progressTextTeams" style="color:var(--text-light); font-size:12px; margin-top:5px; text-align:center;">0% Completed</p>
                    </div>
                    <div id="teamsConnStatus" style="display:none; padding:10px; border-radius:8px; font-size:0.9rem; margin-bottom:10px;"></div>
                </div>

                <!-- Meeting Link Form (shown after connection) -->
                <div id="teamsMeetingForm" style="display:none; border-top:1px solid var(--border); padding-top:15px; margin-top:10px;">
                    <div class="form-group">
                        <label for="teamsMeetingLink">Teams Meeting Link</label>
                        <input type="text" id="teamsMeetingLink" placeholder="https://teams.microsoft.com/l/meetup-join/...">
                    </div>
                    <div class="form-group">
                        <label for="teamsMomTemplate">MOM Format</label>
                        <select id="teamsMomTemplate">
                            <option value="standard">Standard (Concise Summary)</option>
                            <option value="executive">Executive (High-level)</option>
                            <option value="technical">Technical (Engineering)</option>
                            <option value="agile">Agile (Standup)</option>
                            <option value="client">Client Meeting</option>
                        </select>
                    </div>
                    <button id="btnFetchTeamsMOM" onclick="teamsFetchMeeting()">
                        <span class="btn-text">⚡ Fetch Transcript & Generate MOM</span>
                        <div class="spinner"></div>
                    </button>
                    <div id="teamsResult" class="result" style="margin-top:15px;"></div>
                    <button id="btnSaveteams" class="secondary" style="display:none; margin-top:10px;"
                            onclick="saveMOMToFile('teamsResult', 'teams_meeting_mom.txt')">
                        💾 Save MOM
                    </button>
                </div>
            </div>


            <!-- Live Streaming ASR Section -->
            <div class="section">
                <h2>🎤 Live Streaming ASR</h2>
                <p style="margin-top:0; color:var(--text-light); font-size:0.9rem;">Stream real-time audio directly to the server for live transcription using Faster-Whisper.</p>
                
                <div id="recIndicator" class="rec-indicator" style="display:none; color: var(--error); font-weight: bold; margin-bottom: 10px;">
                    <div class="rec-dot" style="display:inline-block; width:10px; height:10px; background:var(--error); border-radius:50%; margin-right:5px; animation: pulse 1s infinite;"></div> Live Streaming...
                </div>

                <div style="display: flex; gap: 10px; align-items: center;">
                    <button id="btnStartLive" onclick="startLiveStreaming()">▶️ Start Live Stream</button>
                    <button id="btnStopLive" class="danger" onclick="stopLiveStreaming()" disabled>⏹️ Stop Stream</button>
                </div>

                <!-- Live Stream Progress Bar — below the buttons -->
                <div id="progressContainerLive" class="progress-container" style="display:none; margin-top:15px;">
                    <div class="progress-bar-bg" style="width:100%; background:#333; border-radius:5px;">
                        <div id="progressBarFillLive" style="width:0%; background:#4CAF50; height:15px; border-radius:5px; transition: width 0.3s;"></div>
                    </div>
                    <p id="progressTextLive" style="color:var(--text-light); font-size:12px; margin-top:5px; text-align:center; margin-bottom:0;">0% Completed</p>
                </div>
                
                <div id="liveTranscript" class="result" style="margin-top: 15px; min-height: 100px; white-space: pre-wrap;"></div>
                
                <div id="exportLiveGrp" style="display:none; margin-top: 10px; gap: 10px; flex-wrap: wrap;">
                    <button class="secondary" onclick="saveMOMToFile('liveTranscript', 'live_meeting_mom.txt')">💾 TXT</button>
                    <button class="secondary" onclick="exportToServer('liveTranscript', 'docx')">📄 DOCX</button>
                    <button class="secondary" onclick="exportToServer('liveTranscript', 'pdf')">📕 PDF</button>
                    <button class="secondary" onclick="exportToServer('liveTranscript', 'template')">📄 Generate MOM Template</button>
                </div>
            </div>
            
            <!-- File Upload Section -->
            <div class="section">
                <h2>📁 Upload Audio/Video File</h2>
                <p style="margin-top:0; color:var(--text-light); font-size:0.85rem;">[OK] Supports: WAV, MP3, M4A, MP4, OGG, AAC, FLAC</p>
                <div class="form-group">
                    <label for="audioFile">Upload Meeting Audio or Video Recording</label>
                    <input type="file" id="audioFile" accept=".wav,.mp3,.m4a,.mp4,.ogg,.aac,.flac">
                    <small style="color: var(--text-light); margin-top: 5px; display: block;">Max file size: Check browser limit. Supported: All audio formats + MP4 video recordings</small>
                </div>
                <div class="form-group">
                    <label for="template">Format Style</label>
                    <select id="template">
                        <option value="standard">Standard (Concise Summary)</option>
                        <option value="executive">Executive (High-level)</option>
                        <option value="agile">Agile (Standup Format)</option>
                        <option value="technical">Technical (Engineering)</option>
                    </select>
                </div>
                <button id="btnAudio" onclick="generateMOM()">
                    <span class="btn-text">Generate MOM</span>
                    <div class="spinner"></div>
                </button>
                
                <!-- Upload Progress Bar -->
                <div id="progressContainerUpload" class="progress-container" style="display:none; margin-top:20px;">
                    <div class="progress-bar-bg" style="width:100%; background:#333; border-radius:5px;">
                        <div id="progressBarFillUpload" style="width:0%; background:#4CAF50; height:15px; border-radius:5px; transition: width 0.3s;"></div>
                    </div>
                    <p id="progressTextUpload" style="color:var(--text-light); font-size:12px; margin-top:5px; text-align:center;">0% Completed</p>
                </div>
                <div id="momResult" class="result"></div>
                <div id="exportUploadGrp" style="display:none; margin-top: 10px; gap: 10px; flex-wrap: wrap;">
                    <button class="secondary" onclick="saveMOMToFile('momResult', 'audio_meeting_mom.txt')">💾 TXT</button>
                    <button class="secondary" onclick="exportToServer('momResult', 'docx')">📄 DOCX</button>
                    <button class="secondary" onclick="exportToServer('momResult', 'pdf')">📕 PDF</button>
                    <button class="secondary" onclick="exportToServer('momResult', 'template')">📄 Generate MOM Template</button>
                </div>
            </div>
            
            <!-- Text Transcript Section -->
            <div class="section">
                <h2>📝 From Text Transcript</h2>
                <div class="form-group">
                    <label for="transcript">Paste Transcript Content</label>
                    <textarea id="transcript" placeholder="Paste your meeting text here..."></textarea>
                </div>
                <div class="form-group">
                    <label for="template2">Format Style</label>
                    <select id="template2">
                        <option value="standard">Standard (Concise Summary)</option>
                        <option value="executive">Executive (High-level)</option>
                        <option value="agile">Agile (Standup Format)</option>
                        <option value="technical">Technical (Engineering)</option>
                    </select>
                </div>
                <button id="btnText" onclick="generateFromText()">
                    <span class="btn-text">Generate MOM</span>
                    <div class="spinner"></div>
                </button>
                <div id="textResult" class="result"></div>
                <button id="btnSavetext" class="secondary" style="display:none; margin-top: 10px;" onclick="saveMOMToFile('textResult', 'text_meeting_mom.txt')">
                    💾 Save MOM
                </button>
            </div>
        </div>
        
        <script>
            // Theme Management
            function toggleTheme() {
                const currentTheme = document.documentElement.getAttribute("data-theme");
                const newTheme = currentTheme === "dark" ? "light" : "dark";
                document.documentElement.setAttribute("data-theme", newTheme);
                localStorage.setItem("selectedTheme", newTheme);
                
                // Keep momSettings in sync
                let settings = JSON.parse(localStorage.getItem("momSettings")) || {};
                settings.theme = newTheme;
                localStorage.setItem("momSettings", JSON.stringify(settings));
            }
            
            // System status check
            window.onload = function() {
                const settings = JSON.parse(localStorage.getItem("momSettings")) || {};
                let savedTheme = settings.theme || localStorage.getItem("selectedTheme");
                
                if (!savedTheme || savedTheme === 'auto') {
                    savedTheme = window.matchMedia("(prefers-color-scheme: dark)").matches ? "dark" : "light";
                }
                document.documentElement.setAttribute("data-theme", savedTheme);

                checkStatus();
                fetchModelConfig();
                
                // Show last generated MOM if redirected from settings
                if (window.location.search.includes('show_last=true')) {
                    const lastMOM = localStorage.getItem('lastGeneratedMOM');
                    const lastTranscript = localStorage.getItem('lastGeneratedTranscript');
                    if (lastMOM) {
                        document.getElementById('transcript').value = lastTranscript || '';
                        showResult('textResult', `<strong>Generated MOM (From Teams):</strong><br><br>${lastMOM.replace(/</g, '&lt;').replace(/>/g, '&gt;').replace(/\\n/g, '<br>')}`);
                        document.getElementById('btnSaveText').style.display = 'block';
                        // Scroll to it
                        setTimeout(() => document.getElementById('textResult').scrollIntoView({behavior: 'smooth'}), 500);
                    }
                }
                setInterval(checkStatus, 30000);
            };
            
            function checkStatus() {
                fetch('/api/v1/status')
                    .then(r => r.json())
                    .then(data => {
                        let html = '';
                        let ollama_ok = data.services.ollama && data.services.ollama.status === 'running';
                        let whisper_ok = data.services.whisper && data.services.whisper.status === 'ready';
                        
                        html += ollama_ok ? '<div class="status-badge status-ok">🟢 AI Ready</div>' : 
                                '<div class="status-badge status-error">🔴 AI Offline</div>';
                        html += whisper_ok ? '<div class="status-badge status-ok">🟢 Whisper Ready</div>' : 
                                '<div class="status-badge status-error">🔴 Whisper Offline</div>';
                        
                        document.getElementById('status').innerHTML = html;
                    }).catch(e => {
                        document.getElementById('status').innerHTML = '<div class="status-badge status-error">🔴 Server Disconnected</div>';
                    });
            }
            
            function fetchModelConfig() {
                fetch('/api/v1/config')
                    .then(r => r.json())
                    .then(data => {
                        if(data.status === 'ok') {
                            const chunking = data.config.chunking_model || 'Not set';
                            const ollama = data.config.ollama_model || 'Not set';
                            const display = document.getElementById('modelConfigDisplay');
                            if(display) {
                                display.innerHTML = `Fast Layer: <strong>${chunking}</strong> &nbsp;|&nbsp; Refinement Layer: <strong>${ollama}</strong>`;
                            }
                        }
                    })
                    .catch(e => console.log('Error fetching config', e));
            }
            
            function setLoading(btnId, isLoading) {
                const btn = document.getElementById(btnId);
                if (!btn) return;
                if (isLoading) {
                    btn.classList.add('btn-loading');
                    btn.disabled = true;
                } else {
                    btn.classList.remove('btn-loading');
                    btn.disabled = false;
                }
            }
            
            function showResult(elementId, content, isError = false, isLoading = false) {
                const el = document.getElementById(elementId);
                el.className = 'result active' + (isError ? ' error' : '');
                
                if (isError) {
                    el.innerHTML = `<strong>Error:</strong> ${content}`;
                } else {
                    el.innerHTML = content.replace(/\\n/g, '<br>');
                }
                
                // Show/hide save button
                const btnId = 'btnSave' + elementId.replace('Result', '');
                const saveBtn = document.getElementById(btnId);
                if (saveBtn) {
                    saveBtn.style.display = (isError || isLoading) ? 'none' : 'inline-block';
                }
            }
            
            function saveMOMToFile(elementId, defaultFilename) {
                const el = document.getElementById(elementId);
                if (!el || !el.innerText) return;
                
                let textContent = el.innerText;
                const marker = "=== FINAL MINUTES OF MEETING ===";
                if(textContent.includes(marker)) {
                    textContent = textContent.split(marker)[1].trim();
                }
                
                const blob = new Blob([textContent], { type: 'text/plain' });
                const url = URL.createObjectURL(blob);
                const a = document.createElement('a');
                a.href = url;
                a.download = defaultFilename;
                document.body.appendChild(a);
                a.click();
                document.body.removeChild(a);
                URL.revokeObjectURL(url);
            }
            
            async function exportToServer(elementId, format) {
                const el = document.getElementById(elementId);
                if (!el) return;
                
                // Prefer the clean stored MOM text (set by stopLiveStreaming)
                // Fall back to innerText extraction if not available
                let textContent = el.dataset.momText || el.innerText;
                
                // Strip the transcript portion if marker exists
                const marker = "=== FINAL MINUTES OF MEETING ===";
                if(textContent.includes(marker)) {
                    textContent = textContent.split(marker)[1].trim();
                }
                
                if (!textContent) { alert('No MOM content found to export.'); return; }
                
                const url = format === 'docx' ? '/api/v1/export/docx' : 
                            format === 'pdf' ? '/api/v1/export/pdf' : 
                            '/api/v1/export/template';
                const defaultFilename = format === 'template' ? 'SECURE_MOM_Template.docx' : 'meeting_mom.' + format;
                
                try {
                    const response = await fetch(url, {
                        method: 'POST',
                        headers: {'Content-Type': 'application/json'},
                        body: JSON.stringify({text: textContent})
                    });
                    if (!response.ok) throw new Error('Export failed');
                    
                    const blob = await response.blob();
                    const downloadUrl = URL.createObjectURL(blob);
                    const a = document.createElement('a');
                    a.href = downloadUrl;
                    a.download = defaultFilename;
                    document.body.appendChild(a);
                    a.click();
                    document.body.removeChild(a);
                    URL.revokeObjectURL(downloadUrl);
                } catch(e) {
                    alert('Error exporting document: ' + e.message);
                }
            }
            
            // --- Live Streaming Logic ---
            let socket = null;
            if (typeof io !== 'undefined') {
                socket = io();
            } else {
                console.warn("Socket.IO not loaded. Retrying...");
            }
            
            function floatTo16BitPCM(input) {
                let output = new Int16Array(input.length);
                for (let i = 0; i < input.length; i++) {
                    let s = Math.max(-1, Math.min(1, input[i]));
                    output[i] = s < 0 ? s * 0x8000 : s * 0x7FFF;
                }
                return output.buffer;
            }

            let audioContext;
            let mediaStream;
            let scriptProcessor;

            async function startLiveStreaming() {
                if (!socket) {
                    alert("Cannot start live stream — Socket.IO failed to load. Please restart the server and try again.");
                    return;
                }
                try {
                    mediaStream = await navigator.mediaDevices.getUserMedia({ audio: true });
                    audioContext = new (window.AudioContext || window.webkitAudioContext)({ sampleRate: 16000 });
                    const source = audioContext.createMediaStreamSource(mediaStream);
                    
                    scriptProcessor = audioContext.createScriptProcessor(4096, 1, 1);
                    
                    scriptProcessor.onaudioprocess = function(e) {
                        const inputData = e.inputBuffer.getChannelData(0);
                        const pcm16 = floatTo16BitPCM(inputData);
                        socket.emit('audio_chunk', pcm16);
                    };
                    
                    source.connect(scriptProcessor);
                    scriptProcessor.connect(audioContext.destination);
                    
                    document.getElementById('btnStartLive').disabled = true;
                    document.getElementById('btnStopLive').disabled = false;
                    document.getElementById('recIndicator').style.display = 'block';
                    
                    const el = document.getElementById('liveTranscript');
                    el.className = 'result active';
                    el.innerHTML = '';
                    
                } catch (err) {
                    alert('Error accessing microphone: ' + err.message);
                }
            }

            async function stopLiveStreaming() {
                if (scriptProcessor) {
                    scriptProcessor.disconnect();
                    scriptProcessor = null;
                }
                if (audioContext && audioContext.state !== 'closed') {
                    audioContext.close();
                }
                if (mediaStream) {
                    mediaStream.getTracks().forEach(track => track.stop());
                }
                document.getElementById('recIndicator').style.display = 'none';
                document.getElementById('btnStartLive').disabled = false;
                document.getElementById('btnStopLive').disabled = true;
                
                // Fetch final MOM
                const el = document.getElementById('liveTranscript');
                el.innerHTML += "\\n\\n[Stream ended. Generating Final MOM...]\\n";
                el.scrollTop = el.scrollHeight;
                
                document.getElementById('progressContainerLive').style.display = 'block';
                document.getElementById('progressBarFillLive').style.width = '0%';
                document.getElementById('progressTextLive').innerText = '0% Completed';
                
                try {
                    const response = await fetch('/api/v1/generate-live-mom', { method: 'POST' });
                    const data = await response.json();
                    if(data.status === 'ok') {
                        // Store the clean MOM text as a data attribute so PDF/DOCX export can read it
                        const momText = data.mom;
                        el.dataset.momText = momText;
                        el.innerHTML += "\\n\\n=== FINAL MINUTES OF MEETING ===\\n\\n" + momText;
                        document.getElementById('exportLiveGrp').style.display = 'flex';
                    } else {
                        el.innerHTML += "\\n\\nError generating MOM: " + data.error;
                    }
                } catch(e) {
                    el.innerHTML += "\\n\\nFailed to generate final MOM: " + e.message;
                }
                el.scrollTop = el.scrollHeight;
            }

            if (socket) {
                socket.on('transcript_update', function(data) {
                    const el = document.getElementById('liveTranscript');
                    el.innerHTML += data.text + " ";
                    // Scroll to bottom
                    el.scrollTop = el.scrollHeight;
                });
                
                // Add unified progress bar listener
                socket.on('progress_update', function(data) {
                    const pct = data.progress;
                    const status = data.status;
                    
                    // Update whichever progress bar is currently visible
                    if (document.getElementById('progressContainerUpload').style.display !== 'none') {
                        document.getElementById('progressBarFillUpload').style.width = pct + '%';
                        document.getElementById('progressTextUpload').innerText = pct + '% - ' + status;
                    } else if (document.getElementById('progressContainerLive').style.display !== 'none') {
                        document.getElementById('progressBarFillLive').style.width = pct + '%';
                        document.getElementById('progressTextLive').innerText = pct + '% - ' + status;
                    } else if (document.getElementById('progressContainerTeams').style.display !== 'none') {
                        document.getElementById('progressBarFillTeams').style.width = pct + '%';
                        document.getElementById('progressTextTeams').innerText = pct + '% - ' + status;
                    }
                });
            }
            
            // --- End Live Streaming Logic ---
            
            // --- Microphone Recording Logic ---
            let recInterval;
            let recSeconds = 0;

            async function startRecording() {
                try {
                    mediaStream = await navigator.mediaDevices.getUserMedia({ audio: true });
                    audioContext = new (window.AudioContext || window.webkitAudioContext)({ sampleRate: 16000 });
                    const source = audioContext.createMediaStreamSource(mediaStream);
                    
                    scriptProcessor = audioContext.createScriptProcessor(4096, 1, 1);
                    audioChunks = [];
                    
                    scriptProcessor.onaudioprocess = function(e) {
                        const inputData = e.inputBuffer.getChannelData(0);
                        audioChunks.push(new Float32Array(inputData));
                    };
                    
                    source.connect(scriptProcessor);
                    scriptProcessor.connect(audioContext.destination);
                    
                    document.getElementById('btnStartRec').disabled = true;
                    document.getElementById('btnStopRec').disabled = false;
                    document.getElementById('recIndicator').classList.add('active');
                    
                    const settings = JSON.parse(localStorage.getItem('momSettings')) || {};
                    const maxRecTime = (settings.audioMaxLength || 30) * 60;
                    
                    recSeconds = 0;
                    recInterval = setInterval(() => {
                        recSeconds++;
                        const m = Math.floor(recSeconds / 60).toString().padStart(2, '0');
                        const s = (recSeconds % 60).toString().padStart(2, '0');
                        document.getElementById('recTime').innerText = `${m}:${s}`;
                        
                        if (recSeconds >= maxRecTime) {
                            stopRecording();
                            alert(`Recording stopped at ${m}:${s} (max time reached)`);
                        }
                    }, 1000);
                    
                } catch (err) {
                    alert('Error accessing microphone: ' + err.message);
                }
            }

            function stopRecording() {
                if (scriptProcessor) {
                    scriptProcessor.disconnect();
                    scriptProcessor = null;
                }
                if (audioContext && audioContext.state !== 'closed') {
                    audioContext.close();
                }
                if (mediaStream) {
                    mediaStream.getTracks().forEach(track => track.stop());
                }
                
                clearInterval(recInterval);
                document.getElementById('recIndicator').classList.remove('active');
                document.getElementById('btnStartRec').disabled = false;
                document.getElementById('btnStopRec').disabled = true;
                
                processRecording();
            }

            function encodeWAV(samples) {
                const buffer = new ArrayBuffer(44 + samples.length * 2);
                const view = new DataView(buffer);
                
                const writeString = function(view, offset, string) {
                    for (let i = 0; i < string.length; i++) {
                        view.setUint8(offset + i, string.charCodeAt(i));
                    }
                };
                
                writeString(view, 0, 'RIFF');
                view.setUint32(4, 36 + samples.length * 2, true);
                writeString(view, 8, 'WAVE');
                writeString(view, 12, 'fmt ');
                view.setUint32(16, 16, true);
                view.setUint16(20, 1, true);
                view.setUint16(22, 1, true);
                view.setUint32(24, 16000, true);
                view.setUint32(28, 16000 * 2, true);
                view.setUint16(32, 2, true);
                view.setUint16(34, 16, true);
                writeString(view, 36, 'data');
                view.setUint32(40, samples.length * 2, true);
                
                let offset = 44;
                for (let i = 0; i < samples.length; i++, offset += 2) {
                    let s = Math.max(-1, Math.min(1, samples[i]));
                    view.setInt16(offset, s < 0 ? s * 0x8000 : s * 0x7FFF, true);
                }
                
                return new Blob([view], { type: 'audio/wav' });
            }

            function processRecording() {
                if (audioChunks.length === 0) return;
                
                let totalLength = audioChunks.reduce((acc, val) => acc + val.length, 0);
                let samples = new Float32Array(totalLength);
                let offset = 0;
                for (let i = 0; i < audioChunks.length; i++) {
                    samples.set(audioChunks[i], offset);
                    offset += audioChunks[i].length;
                }
                
                const audioBlob = encodeWAV(samples);
                const template = document.getElementById('template-mic').value;
                const file = new File([audioBlob], "recorded_meeting.wav", { type: "audio/wav" });
                
                const settings = JSON.parse(localStorage.getItem('momSettings')) || {};
                const wordLimit = settings.momWordLimit || 300;
                
                const formData = new FormData();
                formData.append('audio', file);
                formData.append('template', template);
                formData.append('word_limit', wordLimit);
                
                showResult('micResult', 'Uploading and transcribing recording... This may take a moment.', false, true);
                document.getElementById('btnStartRec').disabled = true;
                
                fetch('/api/v1/generate', { method: 'POST', body: formData })
                .then(r => r.json())
                .then(data => {
                    document.getElementById('btnStartRec').disabled = false;
                    if (data.status === 'success') {
                        const formattedMom = data.mom.replace(/</g, '&lt;').replace(/>/g, '&gt;').replace(/\\n/g, '<br>');
                        showResult('micResult', `<strong>Minutes of Meeting:</strong><br><br>${formattedMom}`);
                    } else {
                        showResult('micResult', data.error, true);
                    }
                })
                .catch(e => {
                    document.getElementById('btnStartRec').disabled = false;
                    showResult('micResult', e.message, true);
                });
            }
            // ---------------------------------
            
            function generateMOM() {
                const file = document.getElementById('audioFile').files[0];
                const template = document.getElementById('template').value;
                
                if (!file) return alert('Please select an audio file');
                
                const settings = JSON.parse(localStorage.getItem('momSettings')) || {};
                const wordLimit = settings.momWordLimit || 300;
                
                const formData = new FormData();
                formData.append('audio', file);
                formData.append('template', template);
                formData.append('word_limit', wordLimit);
                
                setLoading('btnAudio', true);
                
                // Show multi-stage processing
                const stages = [
                    '⏳ Uploading and processing audio...',
                    '⏳ Transcribing with Whisper...',
                    '⏳ Generating structured MOM...',
                    '✨ Finalizing formatting...'
                ];
                let stageIdx = 0;
                showResult('momResult', `<strong>Status:</strong> ${stages[stageIdx]}`, false, true);
                
                const stageInterval = setInterval(() => {
                    stageIdx = Math.min(stageIdx + 1, stages.length - 1);
                    showResult('momResult', `<strong>Status:</strong> ${stages[stageIdx]}`, false, true);
                }, 15000); // Progress stage every 15s to simulate pipeline steps
                
                document.getElementById('progressContainerUpload').style.display = 'block';
                document.getElementById('progressBarFillUpload').style.width = '0%';
                document.getElementById('progressTextUpload').innerText = '0% Completed';
                
                fetch('/api/v1/generate', { method: 'POST', body: formData })
                .then(r => r.json())
                .then(data => {
                    clearInterval(stageInterval);
                    setLoading('btnAudio', false);
                    if (data.status === 'success') {
                        const formattedMom = data.mom.replace(/</g, '&lt;').replace(/>/g, '&gt;').replace(/\\n/g, '<br>');
                        const el = document.getElementById('momResult');
                        el.dataset.momText = data.mom;
                        showResult('momResult', `<strong>Minutes of Meeting:</strong><br><br>${formattedMom}`);
                        document.getElementById('exportUploadGrp').style.display = 'flex';
                    } else {
                        showResult('momResult', data.error, true);
                    }
                })
                .catch(e => {
                    setLoading('btnAudio', false);
                    showResult('momResult', e.message, true);
                });
            }
            
            function generateFromText() {
                const transcript = document.getElementById('transcript').value;
                const template = document.getElementById('template2').value;
                
                if (!transcript.trim()) return alert('Please paste a transcript');
                
                // Check text limit from settings
                const settings = JSON.parse(localStorage.getItem('momSettings')) || {};
                const maxChars = settings.textMaxLength || 10000;
                
                if (transcript.length > maxChars) {
                    alert(`Text exceeds maximum length of ${maxChars} characters. Current: ${transcript.length}`);
                    return;
                }
                
                setLoading('btnText', true);
                showResult('textResult', 'Generating MOM, please wait...', false, true);
                
                fetch('/api/v1/generate-from-text', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({transcript, template, word_limit: settings.momWordLimit || 300})
                })
                .then(r => r.json())
                .then(data => {
                    setLoading('btnText', false);
                    if (data.status === 'success') {
                        const formattedMom = data.mom.replace(/</g, '&lt;').replace(/>/g, '&gt;').replace(/\\n/g, '<br>');
                        showResult('textResult', `<strong>Generated MOM:</strong><br><br>${formattedMom}`);
                    } else {
                        showResult('textResult', data.error, true);
                    }
                })
                .catch(e => {
                    setLoading('btnText', false);
                    showResult('textResult', e.message, true);
                });
            }

            // ─── MS Teams Integration ───────────────────────────────
            let teamsCredentials = null;

            function teamsShowStatus(msg, isError) {
                const el = document.getElementById('teamsConnStatus');
                el.style.display = 'block';
                el.style.background = isError
                    ? 'rgba(239,68,68,0.15)' : 'rgba(16,185,129,0.15)';
                el.style.color = isError ? '#EF4444' : '#10B981';
                el.style.border = `1px solid ${isError ? '#EF4444' : '#10B981'}`;
                el.innerText = msg;
            }

            function teamsGetCreds() {
                return {
                    client_id:     document.getElementById('teamsClientId').value.trim(),
                    api_key:       document.getElementById('teamsClientSecret').value.trim(),
                    tenant_id:     document.getElementById('teamsTenantId').value.trim() || 'common',
                };
            }

            function teamsTestConnection() {
                const creds = teamsGetCreds();
                if (!creds.client_id || !creds.api_key) {
                    teamsShowStatus('Please enter Client ID and Client Secret.', true);
                    return;
                }
                document.getElementById('btnTeamsTest').disabled = true;
                teamsShowStatus('Testing connection...', false);
                fetch('/api/v1/teams/test', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify(creds)
                })
                .then(r => r.json())
                .then(data => {
                    document.getElementById('btnTeamsTest').disabled = false;
                    if (data.status === 'success') {
                        teamsShowStatus('✅ ' + data.message, false);
                    } else {
                        teamsShowStatus('❌ ' + (data.error || 'Connection test failed.'), true);
                    }
                })
                .catch(e => {
                    document.getElementById('btnTeamsTest').disabled = false;
                    teamsShowStatus('❌ ' + e.message, true);
                });
            }

            function teamsConnect() {
                const creds = teamsGetCreds();
                if (!creds.client_id || !creds.api_key) {
                    teamsShowStatus('Please enter Client ID and Client Secret.', true);
                    return;
                }
                document.getElementById('btnTeamsConnect').disabled = true;
                teamsShowStatus('Connecting to Microsoft Teams...', false);
                fetch('/api/v1/teams/test', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify(creds)
                })
                .then(r => r.json())
                .then(data => {
                    document.getElementById('btnTeamsConnect').disabled = false;
                    if (data.status === 'success') {
                        teamsCredentials = creds;
                        teamsShowStatus('✅ Connected! Now paste a meeting link below.', false);
                        document.getElementById('teamsMeetingForm').style.display = 'block';
                    } else {
                        teamsShowStatus('❌ ' + (data.error || 'Authentication failed. Check your credentials.'), true);
                    }
                })
                .catch(e => {
                    document.getElementById('btnTeamsConnect').disabled = false;
                    teamsShowStatus('❌ ' + e.message, true);
                });
            }

            function teamsFetchMeeting() {
                const link = document.getElementById('teamsMeetingLink').value.trim();
                if (!link) { alert('Please paste a Teams meeting link.'); return; }
                if (!teamsCredentials) { alert('Please connect to MS Teams first.'); return; }

                const template   = document.getElementById('teamsMomTemplate').value;
                const settings   = JSON.parse(localStorage.getItem('momSettings')) || {};
                const word_limit = settings.momWordLimit || 300;

                const payload = {
                    ...teamsCredentials,
                    meeting_link: link,
                    template,
                    word_limit,
                };

                setLoading('btnFetchTeamsMOM', true);
                showResult('teamsResult', 'Fetching transcript & generating MOM...', false, true);

                document.getElementById('progressContainerTeams').style.display = 'block';
                document.getElementById('progressBarFillTeams').style.width = '0%';
                document.getElementById('progressTextTeams').innerText = '0% Completed';

                fetch('/api/v1/teams/fetch-meeting', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify(payload)
                })
                .then(r => r.json())
                .then(data => {
                    setLoading('btnFetchTeamsMOM', false);
                    if (data.status === 'success' && data.mom) {
                        const fmt = data.mom.replace(/</g, '&lt;').replace(/>/g, '&gt;').replace(/\\n/g, '<br>');
                        showResult('teamsResult', `<strong>Minutes of Meeting (Teams):</strong><br><br>${fmt}`);
                    } else if (data.warning) {
                        showResult('teamsResult', '⚠️ ' + data.warning + (data.transcript ? '<br><br><em>Transcript fetched but MOM generation failed.</em>' : ''), true);
                    } else {
                        showResult('teamsResult', data.error || 'Failed to fetch meeting.', true);
                    }
                })
                .catch(e => {
                    setLoading('btnFetchTeamsMOM', false);
                    showResult('teamsResult', e.message, true);
                });
            }
            // ────────────────────────────────────────────────────────
        </script>
    </body>
    </html>
    """
    from flask import make_response
    response = make_response(html)
    response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"
    return response


# Error handlers
@app.errorhandler(404)
def not_found(e):
    return jsonify({'error': 'Endpoint not found', 'path': request.path}), 404


@app.errorhandler(500)
def server_error(e):
    logger.error(f"Server error: {str(e)}")
    return jsonify({'error': 'Internal server error'}), 500


def main():
    """Start the server"""
    print("\n" + "="*60)
    print("MOM Generator API Server - Enhanced Edition")
    print("="*60)
    print("\n Running on: http://127.0.0.1:5001")
    print("\n Web Interfaces:")
    print("   - Main UI: http://127.0.0.1:5001/")
    print("   - Settings: http://127.0.0.1:5001/settings")
    print("\n API Documentation:")
    print("   - Health: http://127.0.0.1:5001/health")
    print("   - Status: http://127.0.0.1:5001/api/v1/status")
    print("   - Models: http://127.0.0.1:5001/api/v1/models")
    print("   - Templates: http://127.0.0.1:5001/api/v1/templates")
    print("   - Files: http://127.0.0.1:5001/api/v1/files")
    print("\n Settings Endpoints:")
    print("   - GET /api/v1/settings")
    print("   - POST /api/v1/settings/validate")
    print("\n Upload endpoints:")
    print("   - POST /api/v1/generate (audio/mp4)")
    print("   - POST /api/v1/transcribe (transcribe only)")
    print("   - POST /api/v1/generate-from-text (from text)")
    print("\n MS Teams Integration:")
    print("   - GET /api/v1/teams/status")
    print("   - POST /api/v1/teams/connect")
    print("   - GET /api/v1/teams/list-calls")
    print("   - GET /api/v1/teams/participants/<meeting_id>")
    print("   - POST /api/v1/teams/generate-mom/<meeting_id>")
    print("\n Make sure Ollama is running: ollama serve")
    print("="*60 + "\n")
    
    # Initialize streaming transcriber and pipeline in background
    global streaming_transcriber, streaming_mom_pipeline
    import threading
    
    def init_models_async():
        global generator
        try:
            print("Starting background initialization of AI models...")
            try:
                generator = MOMGenerator(os.path.join(_BASE_DIR, 'config.json'))
                logger.info("MOM Generator initialized successfully in background")
            except Exception as e:
                logger.error(f"Failed to initialize MOM Generator: {str(e)}")
                generator = None
            
            if generator is None:
                return
                
            ollama_client = generator.ollama
            
            # Check Ollama connection and automatically start it if down
            print("Checking Ollama server status and auto-starting if necessary...")
            status = ollama_client.verify_connection()
            if status['connected']:
                print("[OK] Ollama server is running and ready.")
            else:
                print("x Warning: Ollama could not be automatically started. Please open the Ollama app.")
            
            chunking_model = generator.config.get('chunking_model', 'qwen2.5:1.5b')
            refinement_model = generator.config.get('ollama_model', 'qwen2.5:7b')
            
            global streaming_mom_pipeline
            streaming_mom_pipeline = StreamingMOMPipeline(ollama_client, chunking_model, refinement_model)
            streaming_mom_pipeline.start()

            # Load config to get STT model choice
            config = {}
            config_file = os.path.join(_BASE_DIR, "config.json")
            if os.path.exists(config_file):
                with open(config_file, 'r') as f:
                    config = json.load(f)

            stt_model = config.get("stt_model", "distil-small.en")
            global streaming_transcriber
            streaming_transcriber = StreamingTranscriber(device="cpu", model_size=stt_model)
            streaming_transcriber.on_transcription_callback = transcription_callback
            streaming_transcriber.start()
            print("[SUCCESS] All AI models loaded. Ready for transcription!")
        except Exception as e:
            print(f"[ERROR] Failed to load AI models: {e}")

    # Start model loading in a separate thread
    threading.Thread(target=init_models_async, daemon=True).start()

    # Open browser automatically after a short delay
    import webbrowser
    def open_browser():
        webbrowser.open("http://127.0.0.1:5001")
    threading.Timer(1.5, open_browser).start()

    # Run Flask app with socketio immediately
    socketio.run(
        app,
        host='127.0.0.1',
        port=5001,
        debug=False,
        use_reloader=False,
        allow_unsafe_werkzeug=True
    )

if __name__ == '__main__':
    main()
