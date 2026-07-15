import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_template():
    doc = docx.Document()
    
    # 1. Header Table
    table_header = doc.add_table(rows=2, cols=3)
    table_header.style = 'Table Grid'
    # Merge cells for "GENERAL"
    a = table_header.cell(0, 1)
    b = table_header.cell(0, 2)
    
    table_header.cell(0, 0).text = "SECURE"
    table_header.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_header.cell(0, 1).text = "GENERAL\nMinutes Of Meeting"
    table_header.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_header.cell(0, 2).text = "Doc Date   {{ meeting_date }}\nPage       1 of 2"
    
    table_header.cell(1, 0).text = ""
    table_header.cell(1, 1).text = "Minutes Of Meeting"
    table_header.cell(1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_header.cell(1, 2).text = ""
    
    doc.add_paragraph()
    
    # 2. Info Table
    table_info = doc.add_table(rows=6, cols=4)
    table_info.style = 'Table Grid'
    
    def set_row(row_idx, c0, c1, c2, c3):
        table_info.cell(row_idx, 0).text = c0
        table_info.cell(row_idx, 1).text = c1
        table_info.cell(row_idx, 2).text = c2
        table_info.cell(row_idx, 3).text = c3

    set_row(0, "Meeting Date", "{{ meeting_date }}", "Time", "{{ time }}")
    
    # Merge Venue cells
    table_info.cell(1, 0).text = "Venue"
    table_info.cell(1, 1).merge(table_info.cell(1, 3))
    
    # Merge Purpose cells
    table_info.cell(2, 0).text = "Purpose"
    table_info.cell(2, 1).text = "{{ purpose }}"
    table_info.cell(2, 1).merge(table_info.cell(2, 3))
    
    set_row(3, "Planned Duration", "", "Actual Duration", "")
    set_row(4, "Meeting Effort\n(in person hrs.)", "", "Convened By", "")
    set_row(5, "Minutes Taken By", "AI Generator", "Reviewed By", "")

    doc.add_paragraph()

    # 3. Participants
    doc.add_paragraph("1\tParticipants").runs[0].bold = True
    table_part = doc.add_table(rows=2, cols=3)
    table_part.style = 'Table Grid'
    table_part.cell(0, 0).text = "Members Present"
    table_part.cell(0, 1).text = "Members Absent"
    table_part.cell(0, 2).text = "Invitees"
    table_part.cell(1, 0).text = "{{ participants }}"

    doc.add_paragraph()

    # 4. Agenda
    doc.add_paragraph("2\tAgenda").runs[0].bold = True
    table_agenda = doc.add_table(rows=2, cols=2)
    table_agenda.style = 'Table Grid'
    # docxtpl loops need to be placed carefully in table rows
    table_agenda.cell(0, 0).text = "{% tr for item in agenda %}"
    table_agenda.cell(0, 1).text = "" # The row where the loop starts
    
    # Actually, docxtpl loop for rows works by putting the loop start in the first cell of the row
    # and loop end in the last cell of the row, or using {% tr for ... %}
    table_agenda.cell(1, 0).text = "{% tr for item in agenda %}{{ loop.index }}."
    table_agenda.cell(1, 1).text = "{{ item }}{% tr endfor %}"

    doc.add_paragraph()

    # 5. Decisions Summary
    doc.add_paragraph("3\tDecisions Summary").runs[0].bold = True
    table_dec = doc.add_table(rows=2, cols=2)
    table_dec.style = 'Table Grid'
    table_dec.cell(0, 0).text = "Decision No"
    table_dec.cell(0, 1).text = "Decision"
    table_dec.cell(1, 0).text = "{% tr for item in decisions %}{{ loop.index }}"
    table_dec.cell(1, 1).text = "{{ item }}{% tr endfor %}"

    doc.add_paragraph()

    # 6. Discussion Summary
    doc.add_paragraph("4\tDiscussion Summary").runs[0].bold = True
    table_disc = doc.add_table(rows=2, cols=2)
    table_disc.style = 'Table Grid'
    table_disc.cell(0, 0).text = "S.No"
    table_disc.cell(0, 1).text = "Details"
    table_disc.cell(1, 0).text = "{% tr for item in discussions %}{{ loop.index }}"
    table_disc.cell(1, 1).text = "{{ item }}{% tr endfor %}"

    doc.add_paragraph()

    # 7. Action Plan
    doc.add_paragraph("5\tAction Plan").runs[0].bold = True
    table_act = doc.add_table(rows=2, cols=5)
    table_act.style = 'Table Grid'
    table_act.cell(0, 0).text = "Action Item No"
    table_act.cell(0, 1).text = "Action Description"
    table_act.cell(0, 2).text = "Responsibility"
    table_act.cell(0, 3).text = "Target Date"
    table_act.cell(0, 4).text = "Status/Remarks"
    
    table_act.cell(1, 0).text = "{% tr for act in actions %}{{ loop.index }}"
    table_act.cell(1, 1).text = "{{ act.desc }}"
    table_act.cell(1, 2).text = "{{ act.resp }}"
    table_act.cell(1, 3).text = "{{ act.date }}"
    table_act.cell(1, 4).text = "{% tr endfor %}"

    doc.save('source code/backend/SECURE_MOM_Template.docx')
    print("Template SECURE_MOM_Template.docx successfully generated with Jinja tags.")

if __name__ == "__main__":
    create_template()
